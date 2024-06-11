import sqlite3
import pandas as pd
import os
from docx import Document
from docx.shared import Pt, Cm, Mm
import matplotlib.pyplot as plt


database_path = "db_personas.db"

conn = sqlite3.connect("db_personas.db")
query = """
SELECT p.fecha_ingreso AS fecha, s.Rol AS rol, p.residencia AS residencia, p.nombre_completo AS nombre, p.nacionalidad AS nacionalidad, p.fecha_de_nacimiento AS fecha_de_nacimiento, p.profesion AS profesion, s.Sueldo AS salario, p.rut AS rut
FROM personas p
INNER JOIN Salarios s ON p.id_rol = s.id_salarios
"""
df = pd.read_sql_query(query, conn)
conn.close()
print(df)


def get_person_data(df, full_name):
    person_data = df[df['nombre'] == full_name]
    if person_data.empty:
        raise ValueError(f"No se encontró a la persona con nombre: {full_name}")
    return person_data.iloc[0]


def generar_contrato(date, rol, address, rut, full_name, nationality, birth_date, profession, salary):
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Book Antiqua'
    sections = document.sections
    for section in sections:
        section.page_height = Cm(35.56)
        section.page_width = Cm(21.59)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
    
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("header.png")
    document.add_picture('logo.png')
    
    h = document.add_paragraph('')
    h.add_run('CONTRATO DE PRESTACIÓN DE SERVICIOS A HONORARIOS\n').bold = True
    h.alignment = 1
    font.size = Pt(10)
    h.add_run(
        '............................................................................................................................................................................................................................. ').bold = True
    
    p = document.add_paragraph(f'En Temuco, a {str(date)}, entre la Corporación de Innovación y Desarrollo Tecnológico, Rut 78.898.766-4, representada por su Director General don(a) Roberto Gomez Bolainas, Cédula de Identidad Nº 10.678.990-2, ambos domiciliados en Caupolican 455 de esta ciudad, en adelante la “Corporación” y  {full_name}, de nacionalidad {nationality}, de profesión {profession}, nacido el {birth_date}, con domicilio en {address}, Cédula de Identidad N° {rut}, en adelante, el “Prestador de Servicios”, se ha convenido el siguiente contrato de prestación de servicios a honorarios: \n')
    font.size = Pt(8)
    
    p1 = document.add_paragraph('')
    p1.add_run('PRIMERO        :').bold = True
    p1.add_run('En el marco del acuerdo de servicios profesionales fechado el 11 de noviembre de 2020, establecido entre la Agencia Nacional de Estándares Educativos y la Corporación de Innovación y Desarrollo Tecnológico , y ratificado según la Resolución Exenta N°603 del 23 de noviembre de 2020, la Corporación encarga los servicios profesionales del Prestador de Servicios, para que ejecute la siguiente tarea en el proyecto "Evaluación de competencias específicas y metodologías de aprendizaje artificial 2020, ID 67703-20-JJ90."')
    p1.add_run('\n SEGUNDO        :').bold = True
    p1.add_run('El rol a desempeñar es de '+rol+'.')
    p1.add_run('\n TERCERO        :').bold = True
    p1.add_run('El plazo para la realización de la prestación de servicios encomendada será el '+str(date))
    p1.add_run('\n CUARTO        :').bold = True
    p1.add_run('Por el servicio profesional efectivamente realizado, se pagara un monto bruto variable, el cual corresponderá a cada rol dentro de la empresa capacitación, de acuerdo al siguiente detalle: ')
    
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells
    hdr_cells0[0].text = 'Rol'
    hdr_cells0[1].text = 'Monto Bruto'
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = rol
    hdr_cells[1].text = salary
    
    p1.add_run('\n QUINTO        :').bold = True
    p1.add_run('El Prestador de Servicios acepta el encargo y las condiciones precedentes.')
    p1.add_run('\n SEXTO        :').bold = True
    p1.add_run('El Prestador de Servicios está obligado a mantener la confidencialidad de todos los materiales utilizados, conforme al Acuerdo de Confidencialidad previamente establecido.')
    p1.add_run('\n En comprobante, previa lectura y ratificación, las partes firman.  ').bold = True
    
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells[1].add_paragraph()
    r = hdr_cells0.add_run()
    r.add_picture('imagenes/firma.png')
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '-----------------------------------------------------------\nEL PRESTADOR DE SERVICIOS'
    hdr_cells[1].text = '-----------------------------------------------------------\np. LA CORPORACION'
    
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run('Caupolican 0455, Temuco, Chile, www.corpoindet.cl')
    run.add_picture("imagenes/footer1.png")
    
    document.save(f'{full_name}.docx')

def mostrar_menu(df):
    while True:
        print("\nOpciones de generación de contratos:")
        print("1. Generar contratos por selección individual de RUTs (puedes elegir entre un RUT o varios separados por una coma)")
        print("2. Generar contratos por rango de IDs")
        print("3. Salir")
        opcion = input("Seleccione una opción (1, 2 o 3): ")

        if opcion == '1':
            while True:
                    seleccion = input("Ingrese los RUTs de las personas (separados por comas) para generar contratos: ")
                    ruts = [rut.strip() for rut in seleccion.split(',')]
                    personas_encontradas = []
                    for rut in ruts:
                        person_data = df[df['rut'] == rut]
                        if not person_data.empty:
                            personas_encontradas.append(person_data.iloc[0])
                        else:
                            print(f"no se encontro a la persona con RUT:{rut}")

                    if personas_encontradas:
                            for person_data in personas_encontradas:
                                generar_contrato(
                                date=person_data['fecha'],  
                                rol=person_data['rol'],
                                address=person_data['residencia'],  
                                rut=person_data['rut'],  
                                full_name=person_data['nombre'],
                                nationality=person_data['nacionalidad'],
                                birth_date=person_data['fecha_de_nacimiento'],  
                                profession=person_data['profesion'],  
                                salary=str(person_data['salario'])
                            )
                            print(f"Contrato generado para {person_data['nombre']}")
                            break

        elif opcion == '2':
            while True:
                rango = input("Ingrese el rango de IDs para generar contratos (por ejemplo, 0-10): ")
                try:
                    inicio, fin = map(int, rango.split('-'))
                    indices = range(inicio, fin + 1)
                    for i in indices:
                        if i < len(df):
                            person_data = df.iloc[i]
                            generar_contrato(
                                date=person_data["fecha"],  
                                rol=person_data['rol'],
                                address=person_data['residencia'],  
                                rut=person_data['rut'],  
                                full_name=person_data['nombre'],
                                nationality=person_data['nacionalidad'],
                                birth_date=person_data['fecha_de_nacimiento'],  
                                profession=person_data['profesion'],  
                                salary=str(person_data['salario'])
                            )
                            print(f"Contrato generado para {person_data['nombre']}")
                        else:
                            print(f"El ID {i} está fuera del rango de datos disponible.")
                    break 
                except ValueError:
                    print("Rango de IDs inválido. Por favor, intente de nuevo.")
                except Exception as e:
                    print(f"Ha ocurrido un error: {e}")

        elif opcion == '3':
            print("Saliendo del menú.")
            break
        else:
            print("Opción inválida. Por favor, intente de nuevo.")

mostrar_menu(df)




promedio_sueldo = df.groupby('rol')['salario'].mean()
promedio_sueldo.plot(kind='bar')
plt.title('Promedio de Sueldo por Profesión')
plt.xlabel('Profesión')
plt.ylabel('Salario Promedio (CLP)')
plt.show()

distribucion_profesiones = df['rol'].value_counts()
distribucion_profesiones.plot(kind='pie', autopct='%1.1f%%')
plt.title('Distribución de Profesiones')
plt.ylabel('')
plt.show()

conteo_nacionalidades = df['nacionalidad'].value_counts()
conteo_nacionalidades.plot(kind='bar')
plt.title('Conteo de Profesionales por Nacionalidad')
plt.xlabel('Nacionalidad')
plt.ylabel('Cantidad')
plt.show()